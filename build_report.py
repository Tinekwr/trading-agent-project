import os
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
ROOT_DIR = BASE_DIR.parent
RESULTS_DIR = BASE_DIR / "results"
ASSETS_DIR = RESULTS_DIR / "report_assets"
OUT_DOCX = ROOT_DIR / "期末报告456_修订版.docx"


STRATEGY_CN = {
    "Buy and Hold": "买入并持有",
    "Technical Only": "纯技术策略",
    "Technical + Sentiment": "技术+情绪",
    "Long-Biased Multi-Agent": "长期持有偏向多Agent",
    "Tech + Sentiment": "技术+情绪",
}

COLORS = {
    "Buy and Hold": "#4E79A7",
    "Technical Only": "#E15759",
    "Technical + Sentiment": "#59A14F",
    "Long-Biased Multi-Agent": "#F28E2B",
    "Tech + Sentiment": "#59A14F",
}


def font_path(name):
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for item in candidates:
        if item.exists():
            return str(item)
    return None


def pil_font(size, bold=False):
    paths = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for item in paths:
        if item.exists():
            return ImageFont.truetype(str(item), size)
    return ImageFont.load_default()


def setup_matplotlib():
    fp = font_path("cn")
    if fp:
        matplotlib.font_manager.fontManager.addfont(fp)
        prop = matplotlib.font_manager.FontProperties(fname=fp)
        plt.rcParams["font.family"] = prop.get_name()
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.dpi"] = 160


def pct(value):
    return f"{value * 100:.2f}%"


def draw_card(draw, box, title, value, note, fill="#FFFFFF", accent="#2F5597"):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=12, fill=fill, outline="#D8DEE9", width=2)
    draw.rectangle((x1, y1, x1 + 7, y2), fill=accent)
    draw.text((x1 + 22, y1 + 18), title, font=pil_font(24, True), fill="#243447")
    draw.text((x1 + 22, y1 + 54), value, font=pil_font(34, True), fill=accent)
    draw.text((x1 + 22, y1 + 102), note, font=pil_font(20), fill="#5C6670")


def make_frontend_overview(metrics_df):
    out = ASSETS_DIR / "frontend_overview.png"
    img = Image.new("RGB", (1800, 1050), "#F5F7FA")
    draw = ImageDraw.Draw(img)

    draw.rectangle((0, 0, 1800, 88), fill="#FFFFFF")
    draw.text((42, 23), "多Agent量化交易系统 · 回测看板", font=pil_font(36, True), fill="#1D2733")
    draw.text((1245, 31), "AAPL | 2020-2024 | Streamlit", font=pil_font(24), fill="#667085")

    draw.rounded_rectangle((35, 120, 360, 990), radius=18, fill="#FFFFFF", outline="#D8DEE9")
    draw.text((68, 154), "参数设置", font=pil_font(28, True), fill="#1D2733")
    draw.text((68, 212), "回测区间", font=pil_font(22, True), fill="#344054")
    draw.rounded_rectangle((68, 248, 325, 300), radius=8, fill="#F2F4F7", outline="#D0D5DD")
    draw.text((86, 260), "2021-08-01  至  2024-12-31", font=pil_font(18), fill="#344054")
    draw.text((68, 338), "策略显示", font=pil_font(22, True), fill="#344054")
    options = ["买入并持有", "纯技术策略", "技术+情绪", "长期持有偏向多Agent"]
    for i, option in enumerate(options):
        y = 386 + i * 54
        draw.rounded_rectangle((68, y, 102, y + 34), radius=6, fill="#2F5597")
        draw.text((112, y + 2), option, font=pil_font(20), fill="#344054")
    draw.line((68, 635, 325, 635), fill="#E4E7EC", width=2)
    draw.text((68, 675), "Agent 说明", font=pil_font(22, True), fill="#344054")
    notes = [
        "技术Agent：MA / RSI / MACD",
        "情绪Agent：月度新闻评分",
        "风险Agent：风险阈值与回撤保护",
        "决策Agent：长期持有 + 风险过滤",
    ]
    for i, note in enumerate(notes):
        draw.text((68, 722 + i * 44), note, font=pil_font(18), fill="#667085")

    long_row = metrics_df.loc["Long-Biased Multi-Agent"]
    bh_row = metrics_df.loc["Buy and Hold"]
    card_y = 128
    draw_card(draw, (405, card_y, 720, card_y + 150), "累计收益率", pct(long_row["Cumulative Return"]), "多Agent策略表现", accent="#F28E2B")
    draw_card(draw, (748, card_y, 1063, card_y + 150), "夏普比率", f"{long_row['Sharpe Ratio']:.3f}", "风险调整后收益", accent="#2F5597")
    draw_card(draw, (1091, card_y, 1406, card_y + 150), "最大回撤", pct(long_row["Max Drawdown"]), "低于买入并持有", accent="#B42318")
    draw_card(draw, (1434, card_y, 1749, card_y + 150), "对比基准", pct(bh_row["Cumulative Return"]), "Buy and Hold", accent="#667085")

    draw.rounded_rectangle((405, 320, 1749, 990), radius=18, fill="#FFFFFF", outline="#D8DEE9")
    draw.text((440, 350), "核心可视化区域", font=pil_font(28, True), fill="#1D2733")
    draw.text((440, 395), "净值曲线、回撤曲线、技术指标和情绪风险图支持交互式观察，方便从收益、风险和信号来源三个角度解释策略。", font=pil_font(22), fill="#667085")

    chart = Image.open(RESULTS_DIR / "equity_curve.png").convert("RGB")
    chart.thumbnail((1220, 500))
    img.paste(chart, (465, 465))
    draw.text((465, 930), "图中看板使用策略多选和统一悬停提示，适合在答辩时快速切换对比对象。", font=pil_font(22), fill="#475467")
    img.save(out, quality=95)
    return out


def make_frontend_detail():
    out = ASSETS_DIR / "frontend_detail.png"
    img = Image.new("RGB", (1800, 1050), "#F7F8FA")
    draw = ImageDraw.Draw(img)
    draw.text((48, 30), "前端功能分区示意", font=pil_font(38, True), fill="#1D2733")
    draw.text((48, 82), "看板把策略绩效、技术信号、持仓状态、情绪风险和数据明细组织在同一页面中。", font=pil_font(24), fill="#667085")

    panels = [
        ((48, 150, 870, 495), "技术指标走势", "MA20/MA60、RSI、MACD 分层展示，帮助解释技术Agent信号来源。", RESULTS_DIR / "equity_curve.png"),
        ((930, 150, 1752, 495), "回撤监控", "回撤曲线突出风险阶段，便于说明风险Agent的保护作用。", RESULTS_DIR / "drawdown_curve.png"),
        ((48, 560, 870, 970), "情绪与风险评分", "月度 sentiment_score 和 risk_score 展示 LLM 信息进入策略的路径。", RESULTS_DIR / "sentiment_risk_scores.png"),
        ((930, 560, 1752, 970), "绩效汇总表", "指标表集中呈现收益、波动、夏普和最大回撤，支撑结论表达。", RESULTS_DIR / "metrics_comparison.png"),
    ]
    for box, title, desc, path in panels:
        x1, y1, x2, y2 = box
        draw.rounded_rectangle(box, radius=18, fill="#FFFFFF", outline="#D8DEE9", width=2)
        draw.text((x1 + 24, y1 + 20), title, font=pil_font(26, True), fill="#243447")
        draw.text((x1 + 24, y1 + 58), desc, font=pil_font(19), fill="#667085")
        chart = Image.open(path).convert("RGB")
        chart.thumbnail((x2 - x1 - 60, y2 - y1 - 125))
        img.paste(chart, (x1 + 30, y1 + 105))
    img.save(out, quality=95)
    return out


def make_cn_metric_chart(metrics_df):
    out = ASSETS_DIR / "report_metrics_cn.png"
    rows = ["Buy and Hold", "Technical Only", "Technical + Sentiment", "Long-Biased Multi-Agent"]
    fig, axes = plt.subplots(1, 3, figsize=(13.2, 4.6))
    fig.patch.set_facecolor("white")
    specs = [
        ("Cumulative Return", "累计收益率", lambda x: x * 100, "%"),
        ("Sharpe Ratio", "夏普比率", lambda x: x, ""),
        ("Max Drawdown", "最大回撤", lambda x: x * 100, "%"),
    ]
    for ax, (col, title, transform, suffix) in zip(axes, specs):
        values = [transform(metrics_df.loc[r, col]) for r in rows]
        labels = [STRATEGY_CN[r] for r in rows]
        colors = [COLORS[r] for r in rows]
        bars = ax.bar(labels, values, color=colors, width=0.62)
        ax.set_title(title, fontsize=14, fontweight="bold")
        ax.grid(axis="y", alpha=0.28)
        ax.tick_params(axis="x", labelrotation=22)
        for bar, value in zip(bars, values):
            text = f"{value:.2f}{suffix}" if suffix else f"{value:.3f}"
            va = "bottom" if value >= 0 else "top"
            offset = 0.02 * (max(values) - min(values) + 1)
            ax.text(bar.get_x() + bar.get_width() / 2, value + (offset if value >= 0 else -offset), text, ha="center", va=va, fontsize=9)
    fig.suptitle("核心绩效指标对比", fontsize=16, fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.93])
    fig.savefig(out, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return out


def make_cn_annual_chart(annual_df):
    out = ASSETS_DIR / "report_annual_cn.png"
    annual_df = annual_df.set_index("Date")
    cols = ["Buy & Hold", "Technical Only", "Tech + Sentiment", "Long-Biased Multi-Agent"]
    labels = [STRATEGY_CN.get(c, c) for c in cols]
    x = range(len(annual_df.index))
    width = 0.18
    fig, ax = plt.subplots(figsize=(10.8, 5.4))
    for i, col in enumerate(cols):
        values = annual_df[col] * 100
        ax.bar([v + (i - 1.5) * width for v in x], values, width=width, label=labels[i], color=COLORS.get(col, "#888888"))
    ax.axhline(0, color="#344054", linewidth=1)
    ax.set_xticks(list(x))
    ax.set_xticklabels(annual_df.index.astype(str))
    ax.set_ylabel("年度收益率（%）")
    ax.set_title("年度收益对比", fontsize=16, fontweight="bold")
    ax.grid(axis="y", alpha=0.28)
    ax.legend(ncol=2, frameon=False)
    fig.tight_layout()
    fig.savefig(out, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return out


def make_assets():
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    setup_matplotlib()
    metrics_df = pd.read_csv(RESULTS_DIR / "metrics.csv", index_col=0)
    annual_df = pd.read_csv(RESULTS_DIR / "annual_returns.csv")
    assets = {
        "frontend_overview": make_frontend_overview(metrics_df),
        "frontend_detail": make_frontend_detail(),
        "metrics_cn": make_cn_metric_chart(metrics_df),
        "annual_cn": make_cn_annual_chart(annual_df),
    }
    return metrics_df, annual_df, assets


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 15, "1F4E79"),
        ("Heading 2", 13, "2F5597"),
        ("Heading 3", 11.5, "344054"),
    ]:
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_para(doc, text, first_line=True):
    p = doc.add_paragraph()
    if first_line:
        p.paragraph_format.first_line_indent = Pt(21)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Caption"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 89, 89)


def add_picture(doc, path, caption, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)


def add_metrics_table(doc, metrics_df):
    rows = ["Buy and Hold", "Technical Only", "Technical + Sentiment", "Long-Biased Multi-Agent"]
    headers = ["策略", "累计收益率", "年化收益率", "年化波动率", "夏普比率", "最大回撤"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "1F4E79")
        set_cell_text(cell, header, bold=True, color="FFFFFF")
    for row_name in rows:
        row = metrics_df.loc[row_name]
        values = [
            STRATEGY_CN[row_name],
            pct(row["Cumulative Return"]),
            pct(row["Annual Return"]),
            pct(row["Annual Volatility"]),
            f"{row['Sharpe Ratio']:.3f}",
            pct(row["Max Drawdown"]),
        ]
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_text(cell, value)
        if row_name == "Long-Biased Multi-Agent":
            for cell in cells:
                set_cell_shading(cell, "FFF2CC")
    add_caption(doc, "表 1 主要策略绩效指标对比")


def build_doc(metrics_df, annual_df, assets):
    doc = Document()
    set_doc_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("基于多Agent协同的苹果股票量化交易策略研究")
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4E79")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("课程期末报告 · 前端展示与可视化增强版")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string("667085")

    doc.add_paragraph()
    add_para(doc, "摘要：本项目以苹果公司 AAPL 股票为研究对象，使用 Nasdaq 历史日频数据，并结合 DeepSeek 根据月度新闻整理得到的情绪评分和风险评分，构建了一个由技术 Agent、情绪 Agent、风险 Agent 和决策 Agent 协同工作的量化交易系统。系统不仅完成了数据清洗、信号生成、回测和绩效评估，也进一步实现了 Streamlit 前端看板，把策略净值、回撤、技术指标、情绪风险、持仓暴露和绩效表统一展示出来。", False)
    add_para(doc, "从样本内结果看，长期持有偏向的多 Agent 策略在 2021 年 8 月至 2024 年 12 月期间累计收益率为 90.44%，高于买入并持有策略的 67.37%；最大回撤为 -24.25%，低于买入并持有策略的 -31.31%；夏普比率为 0.874，也高于买入并持有策略的 0.622。本文更关注系统搭建、信号分工和可视化解释过程，而不是把单一股票样本的回测结果扩大为普遍结论。", False)
    add_para(doc, "关键词：多Agent；量化交易；DeepSeek；情绪分析；风险控制；Streamlit；可视化看板", False)

    doc.add_heading("一、引言", level=1)
    add_para(doc, "股票市场受到价格趋势、公司新闻、宏观环境、投资者预期和风险偏好等多种因素共同影响。传统量化策略常依赖移动平均线、RSI、MACD 等技术指标，优点是规则清楚、实现简单、可解释性较强，但它们主要来自价格和成交量本身，难以直接吸收新闻文本中包含的情绪和风险信息。")
    add_para(doc, "大语言模型在文本理解和摘要生成方面的能力，为金融文本分析提供了新的工具。与简单情绪词典相比，LLM 可以结合新闻背景、公司状态和市场语境，给出更接近事件含义的情绪判断。本项目没有把 LLM 输出当成每天买卖的高频信号，而是把月度情绪和风险评分作为低频过滤信息，配合技术信号和风险控制共同进入决策。")
    add_para(doc, "本文的重点不是证明某个策略在所有市场中必然有效，而是展示一个较完整的课程项目流程：从数据准备、Agent 分工、回测逻辑，到前端看板和图表解释。这样的组织方式可以让策略表现、风险变化和信号来源更加透明，也便于在答辩或报告阅读中快速理解系统做了哪些工作。")

    doc.add_heading("二、系统总体设计", level=1)
    doc.add_heading("2.1 数据与信号来源", level=2)
    add_para(doc, "项目使用两类数据。第一类是 AAPL 日频交易数据，包括开盘价、最高价、最低价、收盘价和成交量，用于计算收益率、技术指标和回测净值。第二类是月度新闻情绪数据，由 DeepSeek 根据每月相关新闻生成 sentiment_score 和 risk_score，分别表示整体情绪倾向和风险水平。")
    add_para(doc, "由于情绪数据是月度频率，而股票数据是日频频率，项目采用滞后一月的方式对齐数据。例如，1 月新闻总结只用于指导 2 月交易，避免同月后续新闻被提前使用。这个处理虽然降低了信号时效性，但保证了回测中信息使用顺序更合理。")

    doc.add_heading("2.2 Agent 分工", level=2)
    add_bullets(doc, [
        "技术 Agent：基于 MA20/MA60、RSI 和 MACD 判断趋势和动量，输出偏多或偏空的技术信号。",
        "情绪 Agent：使用 DeepSeek 月度 sentiment_score 描述新闻层面的市场态度，主要作为风险过滤信息。",
        "风险 Agent：结合 risk_score 和组合回撤情况识别极端风险，避免策略在风险过高阶段持续暴露。",
        "决策 Agent：将技术、情绪和风险信息合成为最终仓位。最终版本采用长期持有偏向，不做机械投票，只有在趋势和情绪同时转弱或风险过高时才降低仓位。",
    ])
    add_para(doc, "这种设计的好处是模块边界比较清晰。技术指标解释价格走势，情绪评分解释新闻信息，风险模块负责保护机制，决策模块负责把不同频率的信息转成可执行仓位。即使策略表现出现波动，也能回到具体模块分析原因。")

    doc.add_heading("三、前端系统与可视化展示", level=1)
    add_para(doc, "项目使用 Streamlit 和 Plotly 搭建了一个动态回测看板。相比只在终端输出 CSV 或保存几张静态图片，看板把策略指标、净值曲线、回撤曲线、技术信号、持仓状态、月度情绪和风险评分放在同一个交互页面中。用户可以通过左侧侧边栏选择回测日期区间，也可以勾选需要展示的策略，包括 Buy and Hold、Technical Only、Technical + Sentiment 和 Long-Biased Multi-Agent。这样在分析策略时，不需要每次重新运行代码，只要调整前端参数，就能观察不同时间段和不同策略组合下的表现。")
    add_para(doc, "这个前端并不是简单装饰，而是承担了“解释策略”的作用。多 Agent 系统本身包含多个信号来源，如果只给出最终收益，很难看清收益来自哪里、风险在哪个阶段出现、情绪数据是否真的参与了决策。前端看板把这些信息拆开呈现，使技术 Agent、情绪 Agent、风险 Agent 和决策 Agent 的作用都能被观察到。")
    add_picture(doc, assets["frontend_overview"], "图 1 Streamlit 回测看板整体展示", 6.55)
    add_para(doc, "看板顶部使用四个指标卡片展示核心绩效。每张卡片对应一种策略，主指标为累计收益率，并补充夏普比率和最大回撤。这样的设计可以让读者先获得总体印象：买入并持有收益较高但回撤较深，纯技术策略和技术加情绪策略风险较低但收益不足，而长期持有偏向多 Agent 策略在收益和风险调整后表现之间更均衡。")
    add_para(doc, "净值曲线是前端中最重要的展示内容。系统把四类策略的累计净值放在同一坐标系中，并支持鼠标悬停查看同一日期下各策略的净值。通过这张图可以看到，纯技术策略和技术加情绪策略由于较长时间空仓，曲线更平缓，但在 2023 年和 2024 年上涨行情中明显落后；长期持有偏向多 Agent 策略整体更接近长期持有，但在部分下跌阶段能够减少损失，因此最终净值超过买入并持有。")
    add_para(doc, "回撤曲线用于观察策略在不利行情下的压力。前端根据当前选中的净值曲线实时计算阶段高点和相对回撤，并以面积图形式展示。这个图比净值曲线更直接地反映风险体验：如果策略为了降低回撤而长期空仓，回撤自然会小，但也会错过股票上涨带来的主要收益。本文最终选择长期持有偏向，就是在参与上涨行情和控制极端风险之间做折中。")
    add_picture(doc, assets["frontend_detail"], "图 2 前端功能分区与可视化内容", 6.55)
    add_para(doc, "技术信号部分采用三层子图展示：第一层为收盘价、MA20 和 MA60，第二层为 RSI 并标出 70 和 30 两条阈值线，第三层为 MACD 与 MACD Signal。这样的分层方式能把技术 Agent 的判断依据直接展示出来。当短期均线强于长期均线、RSI 未明显超买且 MACD 强于信号线时，技术信号更偏多；当这些条件转弱时，技术 Agent 会给出偏空判断。")
    add_para(doc, "看板还提供持仓分布视角，用堆叠柱状图展示不同策略的持仓天数和空仓天数。这个图可以解释收益差异的来源：Technical Only 和 Technical + Sentiment 空仓比例较高，虽然降低了部分回撤，但也减少了参与上涨行情的机会；Long-Biased Multi-Agent 持仓天数更多，说明最终策略不是频繁择时，而是在长期持有的基础上进行风险过滤。")
    add_para(doc, "情绪 Agent 和风险 Agent 的输出通过月度评分图展示。上半部分用柱状图显示 sentiment_score，正向情绪和负向情绪用不同颜色区分；下半部分用面积图显示 risk_score，并标出 0.75 的风险阈值线。由于情绪数据本身按月更新，看板没有把它伪装成高频日内信号，而是保留月度节奏，提醒使用者低频文本评分更适合做趋势确认和风险过滤。")
    add_para(doc, "在功能组织上，看板把“总览-风险-信号-数据明细”串成一个完整阅读路径。先用指标卡了解结果，再用净值和回撤曲线判断收益路径和风险阶段，随后查看技术指标、持仓分布和情绪风险评分，最后通过绩效表核对具体数值。这种结构比单独堆放图表更适合课程项目展示，也能体现数据处理、回测计算和前端实现之间的衔接。")
    add_bullets(doc, [
        "净值曲线强调收益路径，能直接比较多 Agent 策略与买入并持有、纯技术策略的长期差异。",
        "回撤曲线强调风险暴露，能说明策略在不利行情中是否减少损失。",
        "技术指标图把 MA、RSI、MACD 分层展示，帮助解释技术 Agent 的信号依据。",
        "情绪风险图展示 sentiment_score、risk_score 及风险阈值，说明 LLM 生成数据不是孤立结果，而是进入了策略过滤逻辑。",
        "持仓分布和交易次数图解释策略收益差异的来源：长期持有偏向策略保持较高市场暴露，同时减少频繁调仓。",
    ])

    doc.add_heading("四、回测方法与评价指标", level=1)
    add_para(doc, "本文比较四类策略：买入并持有、纯技术策略、技术加情绪策略、长期持有偏向多 Agent 策略。前三类策略用于提供基准和消融对比，最后一类是项目最终采用的综合策略。评价指标包括累计收益率、年化收益率、年化波动率、夏普比率和最大回撤。")
    add_para(doc, "回测中特别注意收益计算时点。当天收盘后生成的信号不能用于赚取当天收益，因此向量化策略使用 position.shift(1)，循环策略使用上一日仓位 prev_position 计算当日收益。这个处理保证了策略只使用当时已经能够得到的信息。")
    add_metrics_table(doc, metrics_df)

    doc.add_heading("五、实验结果与图表分析", level=1)
    doc.add_heading("5.1 总体绩效", level=2)
    add_picture(doc, assets["metrics_cn"], "图 3 核心绩效指标中文化对比", 6.45)
    add_para(doc, "从核心指标看，长期持有偏向多 Agent 策略的累计收益率为 90.44%，高于买入并持有策略的 67.37%。它的夏普比率为 0.874，高于买入并持有的 0.622，也略高于纯技术策略的 0.828。这个结果说明，在当前 AAPL 样本内，多 Agent 策略不仅保留了主要上涨敞口，也通过风险过滤改善了风险调整后收益。")

    doc.add_heading("5.2 净值曲线与回撤曲线", level=2)
    add_picture(doc, RESULTS_DIR / "equity_curve.png", "图 4 策略累计净值曲线", 6.45)
    add_para(doc, "净值曲线展示了策略收益的累积过程。纯技术和技术加情绪策略由于持仓时间较短，波动较低但收益也明显不足；买入并持有能够充分参与上涨行情，但在下跌阶段承受较大回撤。长期持有偏向多 Agent 策略介于两者之间：保留大部分上涨参与度，同时在部分风险阶段降低暴露。")
    add_picture(doc, RESULTS_DIR / "drawdown_curve.png", "图 5 策略回撤曲线", 6.45)
    add_para(doc, "回撤曲线用于观察策略在不利行情中的损失深度。买入并持有最大回撤为 -31.31%，长期持有偏向多 Agent 策略最大回撤为 -24.25%。纯技术策略和技术加情绪策略回撤更小，但它们牺牲了大量市场暴露，因此需要结合收益和持仓比例一起分析，不能只用单一回撤指标判断策略优劣。")

    doc.add_heading("5.3 年度收益与持仓暴露", level=2)
    add_picture(doc, assets["annual_cn"], "图 6 年度收益对比", 6.35)
    add_para(doc, "年度收益图显示，2022 年是区分策略风险控制效果的关键年份。买入并持有在 2022 年收益为 -26.83%，长期持有偏向多 Agent 策略为 -18.23%，损失相对更小。2023 年和 2024 年市场重新上涨后，多 Agent 策略继续保持较高市场参与度，因此没有像纯技术策略那样明显错过上涨阶段。")
    add_picture(doc, RESULTS_DIR / "exposure_trades.png", "图 7 持仓暴露与调仓次数", 6.35)
    add_para(doc, "持仓暴露图解释了收益差异的来源。纯技术策略平均持仓比例约为 28.08%，技术加情绪策略约为 17.30%，长期持有偏向多 Agent 策略约为 87.91%。同时，多 Agent 策略调仓次数为 22 次，少于纯技术策略的 64 次和技术加情绪策略的 40 次。这说明最终策略不是靠高频切换获利，而是让低频情绪数据以更合适的方式参与风险过滤。")

    doc.add_heading("5.4 情绪与风险数据", level=2)
    add_picture(doc, RESULTS_DIR / "sentiment_risk_scores.png", "图 8 DeepSeek 月度情绪与风险评分", 6.45)
    add_para(doc, "情绪和风险评分体现了新闻文本信息在策略中的作用。sentiment_score 用于描述市场情绪方向，risk_score 用于识别风险较高的月份。由于这些数据按月更新，项目没有把它们设计成每日买卖触发器，而是作为趋势确认和风险过滤条件。这样处理更符合数据频率，也减少了低频文本信号被过度交易化的问题。")

    doc.add_heading("六、项目工作量与实现亮点", level=1)
    add_para(doc, "本项目的工作量主要体现在四个方面。第一，完成了股票数据清洗、技术指标计算、月度情绪数据整理和日频回测数据对齐。第二，按职责拆分了技术、情绪、风险和决策 Agent，使策略逻辑不是简单堆叠信号，而是形成可解释的协作结构。第三，完成了多类策略的回测、指标计算、年度收益分析、持仓暴露分析和图表输出。第四，搭建了 Streamlit 前端看板，使回测结果从静态文件变成可交互的展示系统。")
    add_para(doc, "前端可视化是本项目比较重要的补充。它让策略不只停留在“最终收益是多少”，而是能进一步回答“收益从哪里来”“风险在什么时候出现”“情绪数据如何影响仓位”“为什么最终选择长期持有偏向”。这些问题在课程报告和答辩中都很关键，因为它们能体现项目不仅跑出了结果，也对结果做了结构化解释。")

    doc.add_heading("七、结论与展望", level=1)
    add_para(doc, "本文构建了一个基于技术指标、DeepSeek 月度情绪数据和风险控制机制的多 Agent 量化交易系统，并通过 Streamlit 看板完成了结果展示。在 AAPL 样本区间内，长期持有偏向多 Agent 策略取得了较好的累计收益和夏普比率，同时相对买入并持有降低了最大回撤。这个结果说明，对于长期趋势较明显、但阶段性风险也较突出的股票，把低频情绪风险信息作为过滤器，比把它机械地作为高频买卖信号更合适。")
    add_para(doc, "需要说明的是，当前结果仍主要来自单一股票和样本内回测，不能直接推广到所有股票或未来市场。后续可以从三个方向继续完善：一是扩展到更多股票或 ETF，观察策略在不同资产上的稳定性；二是引入更高频新闻数据，提高情绪信号的时效性；三是让风险 Agent 根据市场波动率动态调整阈值。整体来看，本项目已经完成了从数据、Agent、回测到前端展示的完整闭环，为后续扩展留下了清晰基础。")

    doc.save(OUT_DOCX)
    return OUT_DOCX


def main():
    metrics_df, annual_df, assets = make_assets()
    out = build_doc(metrics_df, annual_df, assets)
    print(out)


if __name__ == "__main__":
    main()
